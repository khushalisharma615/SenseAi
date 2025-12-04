# app.py
import streamlit as st
import spacy
import torch
from sentence_transformers import SentenceTransformer, util
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import numpy as np
import os

# ---------------- Page config & styling ----------------
st.set_page_config(
    page_title="ATS Dashboard — Resume & Cover Letter Generator",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Simple dark-card like styling (works inside Streamlit)
st.markdown(
    """
    <style>
    .stApp {
        background-color: #0e1117;
        color: #e6edf3;
    }
    .card {
        background: linear-gradient(180deg, rgba(255,255,255,0.03), rgba(255,255,255,0.01));
        padding: 1rem;
        border-radius: 12px;
        box-shadow: 0 4px 18px rgba(2,6,23,0.6);
        border: 1px solid rgba(255,255,255,0.03);
    }
    .muted {color: #9aa4b2; font-size: 0.9rem;}
    </style>
    """,
    unsafe_allow_html=True,
)

# ---------------- Load models (cached) ----------------
@st.cache_resource(show_spinner=False)
def load_models():
    nlp = spacy.load("en_core_web_sm")
    sbert = SentenceTransformer("all-mpnet-base-v2")
    return nlp, sbert

nlp, sbert_model = load_models()

def get_timestamp():
    return datetime.now().strftime("%Y%m%d_%H%M%S")

# ---------------- Sidebar (inputs quick access) ----------------
with st.sidebar:
    st.header("Input (Quick)")
    st.markdown("Fill these fields and click **Generate** at the bottom.")
    s_name = st.text_input("Full Name", key="sidebar_name")
    s_email = st.text_input("Email", key="sidebar_email")
    s_mobile = st.text_input("Mobile Number", key="sidebar_mobile")
    s_job_title = st.text_input("Job Title Applying For", key="sidebar_job")
    st.markdown("---")
    st.caption("You can also fill more fields in the main form.")
    st.markdown(
        "<div class='muted'>Built with: spaCy, Sentence-Transformers, scikit-learn, python-docx</div>",
        unsafe_allow_html=True,
    )

# ---------------- Main layout ----------------
st.title("📄 ATS Professional Dashboard")
st.markdown("Dark theme • Metric cards • Progress bars • Downloadable DOCX")

left_col, right_col = st.columns([2, 1])

with left_col:
    st.subheader("Candidate Information")

    # Use existing variables names but bind to Streamlit widgets.
    name = st.text_input("Enter your full name:", value=locals().get('s_name', '') or "")
    email = st.text_input("Enter your email:", value=locals().get('s_email', '') or "")
    mobile = st.text_input("Enter your mobile number:", value=locals().get('s_mobile', '') or "")
    linkedin = st.text_input("Enter your LinkedIn URL:")
    github = st.text_input("Enter your GitHub URL:")
    portfolio = st.text_input("Enter your portfolio URL:")
    skill_set = st.text_input("Enter your acquired skills (comma-separated):")
    professional_summary = st.text_area("Enter your professional summary:")
    education = st.text_area("Enter your education details:")
    college = st.text_input("Enter your college/university name:")
    projects = st.text_input("Enter your key projects (comma-separated):")
    extra_curricular = st.text_input("Enter your extracurricular activities (comma-separated):")
    soft_skills = st.text_input("Enter your soft skills (comma-separated):")
    job_title = st.text_input("Enter the job title you are applying for:", value=locals().get('s_job_title', '') or "")

    st.markdown("---")
    st.subheader("Job Description")
    st.markdown("Paste the full job description below. (You may paste multiple paragraphs.)")
    job_description = st.text_area("Job Description", height=200)

with right_col:
    st.subheader("Summary")
    # Show quick preview cards
    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.metric("Candidate", name or "—")
    st.metric("Applied Role", job_title or "—")
    st.markdown(f"**Contact:** {email or '—'} | {mobile or '—'}")
    st.markdown("</div>", unsafe_allow_html=True)
    st.markdown("")
    st.subheader("Actions")
    st.info("When you click **Generate**, the app will run your original ATS logic and create DOCX files you can download.")

# ---------------- Generate button & processing ----------------
generate = st.button("Generate Resume & Cover Letter", type="primary")

if generate:
    # small progress UI to show activity (keeps logic intact)
    prog = st.progress(0)
    status_text = st.empty()
    status_text.info("Starting processing...")

    # Build resume_text exactly as in original code
    resume_text = f"""
{name}
{professional_summary}
Applying for: {job_title}
Skills: {skill_set}
Education: {education}
College: {college}
Projects: {projects}
Extracurricular Activities: {extra_curricular}
Soft Skills: {soft_skills}
LinkedIn: {linkedin}
GitHub: {github}
Portfolio: {portfolio}
Contact: {email}, {mobile}
"""

    prog.progress(10)
    status_text.info("Extracting skills using spaCy...")

    # ----- Step 1: Extract Skills from Resume -----
    doc = nlp(resume_text)
    skills = list(set([ent.text for ent in doc.ents if ent.label_ in ["ORG", "GPE", "PERSON"]]))
    skills += [chunk.text for chunk in doc.noun_chunks if len(chunk.text) > 2]
    skills = list(set(skills))
    prog.progress(30)

    # Display extracted skills in a neat card
    with st.container():
        st.subheader("🔍 Extracted Skills")
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.write(skills if skills else "No skills extracted.")
        st.markdown("</div>", unsafe_allow_html=True)

    status_text.info("Computing ATS score (SBERT embeddings)...")

    # ----- Step 2: ATS Score -----
    resume_embedding = sbert_model.encode(resume_text, convert_to_tensor=True)
    jd_embedding = sbert_model.encode(job_description, convert_to_tensor=True)
    ats_score = util.pytorch_cos_sim(resume_embedding, jd_embedding).item()
    prog.progress(60)

    # Metric cards and progress bar for ATS score
    ats_percent = round(ats_score * 100, 2)
    c1, c2, c3 = st.columns(3)
    c1.metric("ATS Match", f"{ats_percent} %")
    # progress bar for match
    st.markdown("**Match Progress**")
    st.progress(min(max(int(ats_percent), 0), 100))
    prog.progress(75)

    # ----- Step 3: Cover Letter Generation -----
    cover_letter_body = f"""
I am writing to express my keen interest in the {job_title} position at your esteemed organization. With a strong background in {', '.join(skill_set.split(',')[:3])}, I bring a proven ability to contribute effectively in team-oriented environments while independently handling complex challenges.

During my academic journey at {college}, I have worked on projects such as {projects.split(',')[0].strip()}, which enhanced my practical knowledge and problem-solving skills. My involvement in activities like {extra_curricular.split(',')[0].strip()} has further strengthened my communication and leadership abilities.

This opportunity aligns perfectly with my career goals and passion for growth in {job_title.lower()} roles. I am confident that my technical background and dedication to excellence would make me a strong fit for your team.
"""

    date_today = datetime.today().strftime('%B %d, %Y')
    cover_letter_formatted = f"""
{date_today}

Hiring Manager Name  
Company Name  
Street Address  
City, State Zip Code  

Dear Hiring Manager,

{cover_letter_body.strip()}

Sincerely,  
{name}  
{email} | {mobile}  
LinkedIn: {linkedin}  
GitHub: {github}  
Portfolio: {portfolio}
"""

    # Save Cover Letter as .docx (same filenames as original logic)
    def generate_cover_letter():
        filename = f"cover_letter_{get_timestamp()}.docx"
        doc = Document()
        for line in cover_letter_formatted.strip().split('\n'):
            doc.add_paragraph(line.strip())
        doc.save(filename)
        return filename

    filename_cl = generate_cover_letter()
    prog.progress(85)
    status_text.success("Cover letter generated.")

    # ----- Step 4: Skill Improvement Suggestions -----
    skill_database = ["Python", "NLP", "Deep Learning", "SQL", "TensorFlow", "Data Analysis", "Leadership", "Teamwork"]
    vectorizer = TfidfVectorizer()
    # Keep logic identical: fit_transform on skill_database + skills
    # Be careful if skills is empty — keep same flow
    skill_vectors = vectorizer.fit_transform(skill_database + skills)
    similarities = cosine_similarity(skill_vectors[-len(skills):], skill_vectors[:-len(skills)]) if len(skills) > 0 else np.zeros((0, len(skill_database)))
    missing_skills = [skill_database[idx] for idx in np.argsort(similarities.mean(axis=0))[-3:]] if len(skills) > 0 else skill_database[:3]
    prog.progress(90)

    # Display suggested skills
    st.subheader("💡 Suggested Skills to Learn")
    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.write(missing_skills)
    st.markdown("</div>", unsafe_allow_html=True)

    # ----- Step 5: Resume Generation -----
    def generate_resume():
        filename = f"resume_{get_timestamp()}.docx"
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        doc.add_heading(name, 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"Email: {email} | Mobile: {mobile}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"LinkedIn: {linkedin} | GitHub: {github} | Portfolio: {portfolio}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("\n")

        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(professional_summary)

        doc.add_heading("Skills", level=1)
        for skill in skill_set.split(','):
            doc.add_paragraph(f"- {skill.strip()}", style='List Bullet')

        doc.add_heading("Projects", level=1)
        for proj in projects.split(','):
            doc.add_paragraph(f"- {proj.strip()}", style='List Bullet')

        doc.add_heading("Education", level=1)
        doc.add_paragraph(f"{education}\nCollege/University: {college}")

        doc.add_heading("Extracurricular Activities", level=1)
        for activity in extra_curricular.split(','):
            doc.add_paragraph(f"- {activity.strip()}", style='List Bullet')

        doc.add_heading("Soft Skills", level=1)
        for skill in soft_skills.split(','):
            doc.add_paragraph(f"- {skill.strip()}", style='List Bullet')

        doc.save(filename)
        return filename

    filename_resume = generate_resume()
    prog.progress(97)
    status_text.success("Resume generated.")

    # ----- Step 6: ATS Score Accuracy Level (same function as original) -----
    def calculate_accuracy(predicted_score, threshold=0.75):
        if predicted_score >= threshold:
            return "High Match"
        elif 0.5 <= predicted_score < threshold:
            return "Moderate Match"
        else:
            return "Low Match"

    accuracy_label = calculate_accuracy(ats_score)
    prog.progress(100)
    status_text.success("Processing complete!")

    # ---------------- Show final results in a dashboard layout ----------------
    st.markdown("---")
    st.header("Results Dashboard")

    # Top metrics
    m1, m2, m3 = st.columns(3)
    m1.metric("ATS Score", f"{ats_percent} %")
    m2.metric("Match Level", accuracy_label)
    m3.metric("Files Created", "2 (Resume, Cover Letter)")

    # Download buttons
    dl_col1, dl_col2 = st.columns(2)
    with dl_col1:
        with open(filename_resume, "rb") as f:
            st.download_button("🔽 Download Resume (.docx)", f, file_name=filename_resume)
    with dl_col2:
        with open(filename_cl, "rb") as f:
            st.download_button("🔽 Download Cover Letter (.docx)", f, file_name=filename_cl)

    # Show some details
    st.subheader("ATS Details")
    st.write("**Raw ATS score (0-1):**", ats_score)
    st.write("**Extracted Named Entities and Noun Chunks (used as skills):**")
    st.dataframe({"extracted_skills": skills})

    # cleanup local files? keep them (same behavior as original saved files)
    st.success("All done — files saved on the server and available for download.")

    # small footer
    st.markdown("<div class='muted'>Generated by ATS Dashboard • Keep your input confidential</div>", unsafe_allow_html=True)
