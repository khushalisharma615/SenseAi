# app.py
import streamlit as st
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import numpy as np

# ML imports will be lazy-loaded inside functions to avoid startup crashes
# (sentence-transformers, spacy, torch, sklearn etc.)

st.set_page_config(page_title="SENSAI — Professional Document Suite",
                   layout="wide",
                   initial_sidebar_state="expanded")


# ---------------- Styling ----------------
st.markdown("""
<style>

/*
   ROOT COLORS
    */
:root {
    --primary-text: #D4D4D4;        /* Main text color - VS Code light gray */
    --secondary-text: #9E9E9E;      /* Secondary text / placeholders */
    --accent-blue: #569CD6;         /* For headings, highlights */
    --accent-cyan: #4EC9B0;         /* For subtle highlights */
    --input-bg: #252526;            /* Input field background */
    --sidebar-bg: #1E1E1E;          /* Sidebar background */
    --sidebar-text: #CCCCCC;        /* Sidebar labels text */
}

/* ==============================
   APP BACKGROUND
   ============================== */
html, body, .stApp {
    background-color: #1E1E1E !important;   /* Dark theme base */
    color: var(--primary-text) !important;
    font-family: 'Segoe UI', sans-serif;
}

/* ==============================
   SIDEBAR
   ============================== */
section[data-testid="stSidebar"] > div:first-child {
    background-color: var(--sidebar-bg) !important;
    color: var(--sidebar-text) !important;
    padding: 20px 12px !important;
    border-radius: 0px;  /* remove default dot/border */
}

/* Sidebar input boxes */
section[data-testid="stSidebar"] input,
section[data-testid="stSidebar"] textarea,
section[data-testid="stSidebar"] select {
    background-color: var(--input-bg) !important;
    color: var(--primary-text) !important;
    border: 1px solid #3C3C3C !important;
    border-radius: 6px !important;
    padding: 6px 10px !important;
}

/* Sidebar labels & text */
section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] span,
section[data-testid="stSidebar"] div {
    color: var(--sidebar-text) !important;
}
section[data-testid="stSidebar"] .stSelectbox > div svg {
    fill: #569CD6 !important;  
    width: 14px !important;
    height: 14px !important;
}

/* Hover effect for selectbox */
section[data-testid="stSidebar"] .stSelectbox > div:hover {
    background-color: #2A2A2A !important;
    cursor: pointer;
}
.stExpander, .stExpander > div {
    background-color: #1E1E1E !important;  /* dark gray */
    color: #D4D4D4 !important;
}
.stExpander summary {
    color: #D4D4D4 !important;
}

/* Remove extra dark overlay when dropdown is opened */
section[data-testid="stSidebar"] div[role="listbox"] {
    background-color: #1E1E1E !important;
    border: 1px solid #3C3C3C !important;
    border-radius: 6px !important;
}
/* ==============================
   MAIN HEADINGS
   ============================== */
.stApp .main-heading {
    font-size: 56px;
    font-weight: 900;
    text-align: center;
    color: #40E0D0 !important;  /* turquoise */
    margin-top: 0px;
}

.sub-heading {
    font-size: 22px;
    font-weight: 500;
    text-align: center;
    color: var(--accent-blue) !important; /* secondary color for subheading */
    margin-top: 0px;
    margin-bottom: 20px;
}

/* ==============================
   MAIN CONTENT TEXT
   ============================== */
.stApp label, 
.stApp p, 
.stApp span, 
.stApp div {
    color: var(--primary-text) !important;
}

/* Inputs in main content */
.stTextInput input,
.stTextArea textarea,
.stSelectbox select {
    background-color: var(--input-bg) !important;
    color: #569CD6 !important;  /* blue text */
    border: 1px solid #3C3C3C !important;
    border-radius: 6px !important;
    padding: 6px 10px !important;
}

/* ==============================
   PREVIEW CARD
   ============================== */
.card, .card * {
    color: var(--primary-text) !important;
    background-color: #2A2A2A !important;
    border-radius: 6px;
    padding: 4px 8px !important;
}

/* Section titles (Candidate Information, Preview, etc.) */
.stHeader, .stSubheader, h1, h2, h3, .stMarkdown h2, .stMarkdown h3 {
    color: #FFFFFF !important;  /* white for section headings */
}

/* Subtle scrollbars for dark theme */
::-webkit-scrollbar {
    width: 8px;
}
::-webkit-scrollbar-track {
    background: #1E1E1E;
}
::-webkit-scrollbar-thumb {
    background-color: #555;
    border-radius: 10px;
    border: 2px solid #1E1E1E;
}
/* ==============================
   Sidebar model select dropdown fix
   ============================== */
section[data-testid="stSidebar"] .stSelectbox > div svg {
    fill: #569CD6 !important;   /* blue arrow */
}

section[data-testid="stSidebar"] div[role="listbox"] {
    background-color: #1E1E1E !important;  /* dark background for dropdown options */
    border: 1px solid #3C3C3C !important;
    border-radius: 6px !important;
}

/* ==============================
   Override specificity for headings
   ============================== */
.stApp .main-heading {
    color: #40E0D0 !important;  /* turquoise main heading */
}

.stApp .stHeader,
.stApp .stSubheader {
    color: #FFFFFF !important;   /* white for Candidate Info, Preview, etc. */
}

</style>
""", unsafe_allow_html=True)

# ------------------------- DISPLAY HEADINGS -------------------------
st.markdown('<div class="main-heading">SENSAI</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-heading">Professional Document Suite</div>', unsafe_allow_html=True) 


# ---------------- Helpers ----------------
def get_timestamp():
    return datetime.now().strftime("%Y%m%d_%H%M%S")


# Lazy model loader with caching and graceful fallback
@st.cache_resource(show_spinner=False)
def load_models(model_choice: str = "all-mpnet-base-v2"):
    """
    Loads spaCy and SentenceTransformer models.
    - model_choice can be a smaller model to reduce memory/time.
    Returns (nlp, sbert_model, util_module) or raises informative Exception.
    """
    # Import heavy libs here to avoid import-time crashes in Streamlit before user action
    try:
        import spacy
    except Exception as e:
        raise RuntimeError("spaCy import failed. Make sure 'spacy' is installed.") from e

    # Ensure spaCy model exists, try to download if missing
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        # attempt programmatic download (best-effort)
        try:
            st.info("Downloading spaCy model 'en_core_web_sm' (one-time). This may take a moment...")
            import spacy.cli
            spacy.cli.download("en_core_web_sm")
            nlp = spacy.load("en_core_web_sm")
        except Exception as e:
            raise RuntimeError(
                "Could not load or download spaCy model 'en_core_web_sm'. "
                "Run 'python -m spacy download en_core_web_sm' manually."
            ) from e

    # SentenceTransformers + useful util
    try:
        from sentence_transformers import SentenceTransformer, util
    except Exception as e:
        raise RuntimeError(
            "Could not import sentence-transformers. Ensure correct versions: "
            "sentence-transformers==2.2.2 and huggingface_hub==0.16.4 (see requirements)."
        ) from e

    # Try to load the chosen SBERT model. If it fails (OOM / network), fall back to a smaller model.
    try:
        sbert = SentenceTransformer(model_choice)
    except Exception as e:
        # fallback to smaller model
        fallback = "all-MiniLM-L6-v2"
        try:
            st.warning(f"Failed to load '{model_choice}', falling back to '{fallback}'.")
            sbert = SentenceTransformer(fallback)
        except Exception as ee:
            raise RuntimeError(
                f"Failed to load SBERT models ('{model_choice}' and fallback '{fallback}'). "
                "This can happen due to network or memory limits."
            ) from ee

    return nlp, sbert, util


# ---------------- Sidebar quick inputs ----------------
with st.sidebar:
    st.header("Quick Inputs")
    st.markdown("Fill fields below & press **Generate**.")
    quick_name = st.text_input("Full Name (quick)", key="quick_name")
    quick_email = st.text_input("Email (quick)", key="quick_email")
    quick_mobile = st.text_input("Mobile (quick)", key="quick_mobile")
    st.markdown("---")
    st.subheader("Model")
    model_choice = st.selectbox(
        "SBERT model (smaller = faster, lighter)",
        options=["all-mpnet-base-v2", "all-MiniLM-L6-v2"],
        help="Larger models give better embeddings but need more memory and time."
    )
    st.markdown("<div class='muted'>Recommended Python: 3.11 • Use CPU-friendly torch if needed</div>", unsafe_allow_html=True)


# ---------------- Main form ----------------

left, right = st.columns([2,1])

with left:
    st.subheader("Candidate Information")
    name = st.text_input("Enter your full name:", value=quick_name or "")
    email = st.text_input("Enter your email:", value=quick_email or "")
    mobile = st.text_input("Enter your mobile number:", value=quick_mobile or "")
    linkedin = st.text_input("Enter your LinkedIn URL:")
    github = st.text_input("Enter your GitHub URL:")
    portfolio = st.text_input("Enter your portfolio URL:")
    skill_set = st.text_input("Enter your acquired skills (comma-separated):")
    professional_summary = st.text_area("Enter your professional summary:")
    education = st.text_area("Enter your education details:")
    college = st.text_input("Enter your college/university name:")
    projects = st.text_input("Enter your key projects (comma-separated):")
    extra_curricular = st.text_input("Enter your extracurricular activities (comma-separated):")
    soft_skills = st.text_input("Enter your soft skills (comma-separated):")
    job_title = st.text_input("Enter the job title you are applying for:")

    st.markdown("---")
    st.subheader("Job Description")
    job_description = st.text_area("Paste the job description (optional)", height=200)

with right:
    st.subheader("Preview")
    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.markdown(f"**Candidate:** {name or '—'}  ")
    st.markdown(f"**Role:** {job_title or '—'}  ")
    st.markdown(f"**Contact:** {email or '—'} • {mobile or '—'}  ")
    st.markdown("</div>", unsafe_allow_html=True)
    st.markdown("")
    st.info("Press **Generate** to run the exact ATS/resume logic and produce DOCX files.")


# ---------------- Generate action ----------------
if st.button("Generate Resume & Cover Letter"):
    status = st.empty()
    progress = st.progress(0)
    try:
        status.info("Loading models (this may take 10–60s depending on model choice & network)...")
        nlp, sbert_model, sbert_util = load_models(model_choice)
    except Exception as e:
        st.error(f"Model loading failed: {e}")
        raise st.stop()

    progress.progress(10)
    status.info("Building resume text...")

    # Build resume_text exactly like original
    resume_text = f"""
{name}
{professional_summary}
Applying for: {job_title}
Skills: {skill_set}
Education: {education}
College: {college}
Projects: {projects}
Extracurricular Activities: {extra_curricular}
Soft Skills: {soft_skills}
LinkedIn: {linkedin}
GitHub: {github}
Portfolio: {portfolio}
Contact: {email}, {mobile}
"""

    progress.progress(25)
    status.info("Extracting skills using spaCy...")

    # ----- Step 1: Extract Skills from Resume -----
    doc = nlp(resume_text)
    # keep original named-entity filtering as you used
    skills = list(set([ent.text for ent in doc.ents if ent.label_ in ["ORG", "GPE", "PERSON"]]))
    skills += [chunk.text for chunk in doc.noun_chunks if len(chunk.text) > 2]
    skills = list(set(skills))

    progress.progress(40)
    status.info("Computing ATS score (SBERT embeddings)...")

    # ----- Step 2: ATS Score -----
    # handle empty job_description safely
    jd_text = job_description or ""
    try:
        resume_embedding = sbert_model.encode(resume_text, convert_to_tensor=True)
        jd_embedding = sbert_model.encode(jd_text, convert_to_tensor=True)
        ats_score = float(sbert_util.pytorch_cos_sim(resume_embedding, jd_embedding).item())
    except Exception as e:
        # if model encoding fails, set safe default and show error
        ats_score = 0.0
        st.warning(f"Embedding computation failed: {e}. ATS score set to 0.0")

    progress.progress(60)
    ats_percent = round(ats_score * 100, 2)

    # show extracted skills
    st.subheader("🔍 Extracted Skills")
    st.markdown("<div class='card'>", unsafe_allow_html=True)
    if skills:
        st.write(skills)
    else:
        st.write("No named entities/noun chunks extracted — consider adding more details in 'Professional Summary'.")
    st.markdown("</div>", unsafe_allow_html=True)

    # metric and progress
    st.markdown("### ATS Match")
    col1, col2 = st.columns([1,3])
    col1.metric("ATS Match", f"{ats_percent} %")
    col2.progress(min(max(int(ats_percent), 0), 100))
    progress.progress(75)

    # ----- Step 3: Cover Letter Generation -----
    cover_letter_body = f"""
I am writing to express my keen interest in the {job_title} position at your esteemed organization. With a strong background in {', '.join(skill_set.split(',')[:3])}, I bring a proven ability to contribute effectively in team-oriented environments while independently handling complex challenges.

During my academic journey at {college}, I have worked on projects such as {projects.split(',')[0].strip() if projects else ''}, which enhanced my practical knowledge and problem-solving skills. My involvement in activities like {extra_curricular.split(',')[0].strip() if extra_curricular else ''} has further strengthened my communication and leadership abilities.

This opportunity aligns perfectly with my career goals and passion for growth in {job_title.lower() if job_title else ''} roles. I am confident that my technical background and dedication to excellence would make me a strong fit for your team.
"""

    date_today = datetime.today().strftime('%B %d, %Y')
    cover_letter_formatted = f"""
{date_today}

Hiring Manager Name  
Company Name  
Street Address  
City, State Zip Code  

Dear Hiring Manager,

{cover_letter_body.strip()}

Sincerely,  
{name or ''}  
{email or ''} | {mobile or ''}  
LinkedIn: {linkedin or ''}  
GitHub: {github or ''}  
Portfolio: {portfolio or ''}
"""

    # Save Cover Letter
    def generate_cover_letter_file():
        filename = f"cover_letter_{get_timestamp()}.docx"
        doc = Document()
        for line in cover_letter_formatted.strip().split('\n'):
            doc.add_paragraph(line.strip())
        doc.save(filename)
        return filename

    filename_cl = generate_cover_letter_file()
    progress.progress(85)
    status.success("Cover letter generated.")

    # ----- Step 4: Skill Improvement Suggestions -----
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity

    skill_database = ["Python", "NLP", "Deep Learning", "SQL", "TensorFlow", "Data Analysis", "Leadership", "Teamwork"]
    vectorizer = TfidfVectorizer()

    # ensure safe for empty skills
    try:
        if skills:
            skill_vectors = vectorizer.fit_transform(skill_database + skills)
            # compute similarities from user skills to skill_database
            similarities = cosine_similarity(skill_vectors[-len(skills):], skill_vectors[:-len(skills)])
            missing_skills = [skill_database[idx] for idx in np.argsort(similarities.mean(axis=0))[-3:]]
        else:
            missing_skills = skill_database[:3]
    except Exception as e:
        missing_skills = skill_database[:3]
        st.warning(f"Skill suggestion step failed: {e} — showing default suggestions.")

    progress.progress(90)

    st.subheader("💡 Suggested Skills to Learn")
    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.write(missing_skills)
    st.markdown("</div>", unsafe_allow_html=True)

    # ----- Step 5: Resume Generation (keeps your original formatting & behavior) -----
    def generate_resume_file():
        filename = f"resume_{get_timestamp()}.docx"
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        doc.add_heading(name or "", 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"Email: {email or ''} | Mobile: {mobile or ''}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"LinkedIn: {linkedin or ''} | GitHub: {github or ''} | Portfolio: {portfolio or ''}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("\n")

        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(professional_summary or "")

        doc.add_heading("Skills", level=1)
        for skill in (skill_set.split(",") if skill_set else []):
            doc.add_paragraph(f"- {skill.strip()}", style='List Bullet')

        doc.add_heading("Projects", level=1)
        for proj in (projects.split(",") if projects else []):
            doc.add_paragraph(f"- {proj.strip()}", style='List Bullet')

        doc.add_heading("Education", level=1)
        doc.add_paragraph(f"{education or ''}\nCollege/University: {college or ''}")

        doc.add_heading("Extracurricular Activities", level=1)
        for activity in (extra_curricular.split(",") if extra_curricular else []):
            doc.add_paragraph(f"- {activity.strip()}", style='List Bullet')

        doc.add_heading("Soft Skills", level=1)
        for skill in (soft_skills.split(",") if soft_skills else []):
            doc.add_paragraph(f"- {skill.strip()}", style='List Bullet')

        doc.save(filename)
        return filename

    filename_resume = generate_resume_file()
    progress.progress(97)
    status.success("Resume generated.")

    # ----- Score label (same function) -----
    def calculate_accuracy(predicted_score, threshold=0.75):
        if predicted_score >= threshold:
            return "High Match"
        elif 0.5 <= predicted_score < threshold:
            return "Moderate Match"
        else:
            return "Low Match"

    accuracy_label = calculate_accuracy(ats_score)

    # final dashboard
    st.markdown("---")
    st.header("Results Dashboard")
    m1, m2, m3 = st.columns(3)
    m1.metric("ATS Score", f"{ats_percent} %")
    m2.metric("Match Level", accuracy_label)
    m3.metric("Files Created", "2")

    dl1, dl2 = st.columns(2)
    with dl1:
        with open(filename_resume, "rb") as f:
            st.download_button("🔽 Download Resume (.docx)", f, file_name=filename_resume)
    with dl2:
        with open(filename_cl, "rb") as f:
            st.download_button("🔽 Download Cover Letter (.docx)", f, file_name=filename_cl)

    st.subheader("ATS Details")
    st.write("**Raw ATS score (0-1):**", ats_score)
    st.dataframe({"extracted_skills": skills})

    progress.progress(100)
    st.success("Done — files saved locally and ready for download.")
    st.markdown("<div class='muted'>Note: Large SBERT models can take time or fail on low-memory hosts (Streamlit Cloud free tier).</div>", unsafe_allow_html=True)
