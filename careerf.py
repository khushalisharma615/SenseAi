import spacy
import torch
from sentence_transformers import SentenceTransformer, util
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import numpy as np

# Load NLP models
nlp = spacy.load("en_core_web_sm")
sbert_model = SentenceTransformer("all-mpnet-base-v2")

# Function to get current timestamp for file naming
def get_timestamp():
    return datetime.now().strftime("%Y%m%d_%H%M%S")

# User Input
name = input("Enter your full name: ")
email = input("Enter your email: ")
mobile = input("Enter your mobile number: ")
linkedin = input("Enter your LinkedIn URL: ")
github = input("Enter your GitHub URL: ")
portfolio = input("Enter your portfolio URL: ")
skill_set = input("Enter your acquired skills (comma-separated): ")
professional_summary = input("Enter your professional summary: ")
education = input("Enter your education details: ")
college = input("Enter your college/university name: ")
projects = input("Enter your key projects (comma-separated): ")
extra_curricular = input("Enter your extracurricular activities (comma-separated): ")
soft_skills = input("Enter your soft skills (comma-separated): ")
job_title = input("Enter the job title you are applying for: ")

# Improved Job Description Input
def get_job_description():
    print("\nPaste the job description (type 'END' on a new line when finished):")
    jd_lines = []
    while True:
        line = input("> ")
        if line.strip().upper() == "END":
            break
        elif line.strip() == "":
            continue  # skip accidental blank lines
        jd_lines.append(line.strip())
    return " ".join(jd_lines)

job_description = get_job_description()
# Resume Text
resume_text = f"""
{name}
{professional_summary}
Applying for: {job_title}
Skills: {skill_set}
Education: {education}
College: {college}
Projects: {projects}
Extracurricular Activities: {extra_curricular}
Soft Skills: {soft_skills}
LinkedIn: {linkedin}
GitHub: {github}
Portfolio: {portfolio}
Contact: {email}, {mobile}
"""

# ----- Step 1: Extract Skills from Resume -----
doc = nlp(resume_text)
skills = list(set([ent.text for ent in doc.ents if ent.label_ in ["ORG", "GPE", "PERSON"]]))
skills += [chunk.text for chunk in doc.noun_chunks if len(chunk.text) > 2]
skills = list(set(skills))
print(f"Extracted Skills: {skills}")

# ----- Step 2: ATS Score -----
resume_embedding = sbert_model.encode(resume_text, convert_to_tensor=True)
jd_embedding = sbert_model.encode(job_description, convert_to_tensor=True)
ats_score = util.pytorch_cos_sim(resume_embedding, jd_embedding).item()
print(f"ATS Match Score: {round(ats_score * 100, 2)}%")

# ----- Step 3: Cover Letter Generation -----
cover_letter_body = f"""
I am writing to express my keen interest in the {job_title} position at your esteemed organization. With a strong background in {', '.join(skill_set.split(',')[:3])}, I bring a proven ability to contribute effectively in team-oriented environments while independently handling complex challenges.

During my academic journey at {college}, I have worked on projects such as {projects.split(',')[0].strip()}, which enhanced my practical knowledge and problem-solving skills. My involvement in activities like {extra_curricular.split(',')[0].strip()} has further strengthened my communication and leadership abilities.

This opportunity aligns perfectly with my career goals and passion for growth in {job_title.lower()} roles. I am confident that my technical background and dedication to excellence would make me a strong fit for your team.
"""

date_today = datetime.today().strftime('%B %d, %Y')
cover_letter_formatted = f"""
{date_today}

Hiring Manager Name  
Company Name  
Street Address  
City, State Zip Code  

Dear Hiring Manager,

{cover_letter_body.strip()}

Sincerely,  
{name}  
{email} | {mobile}  
LinkedIn: {linkedin}  
GitHub: {github}  
Portfolio: {portfolio}
"""

# Save Cover Letter as .docx
def generate_cover_letter():
    filename = f"cover_letter_{get_timestamp()}.docx"
    doc = Document()
    for line in cover_letter_formatted.strip().split('\n'):
        doc.add_paragraph(line.strip())
    doc.save(filename)
    print(f"Cover letter saved as '{filename}'")

generate_cover_letter()

# ----- Step 4: Skill Improvement Suggestions -----
skill_database = ["Python", "NLP", "Deep Learning", "SQL", "TensorFlow", "Data Analysis", "Leadership", "Teamwork"]
vectorizer = TfidfVectorizer()
skill_vectors = vectorizer.fit_transform(skill_database + skills)
similarities = cosine_similarity(skill_vectors[-len(skills):], skill_vectors[:-len(skills)])
missing_skills = [skill_database[idx] for idx in np.argsort(similarities.mean(axis=0))[-3:]]
print(f"Suggested Skills to Learn: {missing_skills}")

# ----- Step 5: Resume Generation -----
def generate_resume():
    filename = f"resume_{get_timestamp()}.docx"
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    doc.add_heading(name, 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"Email: {email} | Mobile: {mobile}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"LinkedIn: {linkedin} | GitHub: {github} | Portfolio: {portfolio}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("\n")

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(professional_summary)

    doc.add_heading("Skills", level=1)
    for skill in skill_set.split(','):
        doc.add_paragraph(f"- {skill.strip()}", style='List Bullet')

    doc.add_heading("Projects", level=1)
    for proj in projects.split(','):
        doc.add_paragraph(f"- {proj.strip()}", style='List Bullet')

    doc.add_heading("Education", level=1)
    doc.add_paragraph(f"{education}\nCollege/University: {college}")

    doc.add_heading("Extracurricular Activities", level=1)
    for activity in extra_curricular.split(','):
        doc.add_paragraph(f"- {activity.strip()}", style='List Bullet')

    doc.add_heading("Soft Skills", level=1)
    for skill in soft_skills.split(','):
        doc.add_paragraph(f"- {skill.strip()}", style='List Bullet')

    doc.save(filename)
    print(f"Resume saved as '{filename}'")

generate_resume()

# ----- Step 6: ATS Score Accuracy Level -----
def calculate_accuracy(predicted_score, threshold=0.75):
    if predicted_score >= threshold:
        return "High Match"
    elif 0.5 <= predicted_score < threshold:
        return "Moderate Match"
    else:
        return "Low Match"

accuracy_label = calculate_accuracy(ats_score)
print(f"ATS Score Accuracy: {accuracy_label}")
